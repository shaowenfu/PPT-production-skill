from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from PIL import Image

from scripts.docx_to_text import _render_docx_to_text


def _make_png_bytes() -> bytes:
    image = Image.new("RGB", (120, 60), color=(20, 20, 20))
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def test_render_docx_to_text_keeps_paragraphs_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("第一页标题")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "主讲"
    table.cell(1, 1).text = "李达聪"
    document.save(docx_path)

    rendered, metrics, artifacts = _render_docx_to_text(docx_path, tmp_path / "pic")

    assert "第一页标题" in rendered
    assert "[表格]" in rendered
    assert "字段 | 内容" in rendered
    assert "主讲 | 李达聪" in rendered
    assert metrics["paragraphs"] == 1
    assert metrics["tables"] == 1
    assert artifacts == []


def test_render_docx_to_text_extracts_images_and_inserts_placeholder(tmp_path: Path) -> None:
    docx_path = tmp_path / "image.docx"
    image_path = tmp_path / "embedded.png"
    image_path.write_bytes(_make_png_bytes())

    document = Document()
    document.add_paragraph("含图片页")
    document.add_picture(str(image_path))
    document.save(docx_path)

    rendered, metrics, artifacts = _render_docx_to_text(docx_path, tmp_path / "pic")

    assert "含图片页" in rendered
    assert "[此处有图片]" in rendered
    assert metrics["images"] == 1
    assert len(artifacts) == 1
    assert Path(artifacts[0]).exists()
    assert artifacts[0] in rendered
